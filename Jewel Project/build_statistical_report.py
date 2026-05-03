from __future__ import annotations

from datetime import date
from pathlib import Path
import math
import textwrap

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import sys

DOC_SKILL_DIR = Path("/Users/freshliannes.rosal/.codex/plugins/cache/openai-primary-runtime/documents/26.430.10722/skills/documents")
sys.path.append(str(DOC_SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights


ROOT = Path(".")
OUT = ROOT / "statistical_report_outputs"
FIG = OUT / "figures"
OUT.mkdir(exist_ok=True)
FIG.mkdir(exist_ok=True)

REPORT_DOCX = OUT / "Statistical_Report_Poso_Bacterial_Profile_Engr_Jamie_Eduardo_Rosal.docx"
CONTENT_WIDTH_DXA = 9936


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
    ]
    for path in candidates:
        if path and Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


FONT_TITLE = load_font(34, True)
FONT_SUBTITLE = load_font(22, True)
FONT_LABEL = load_font(20)
FONT_SMALL = load_font(16)
FONT_TINY = load_font(13)
FONT_BOLD = load_font(18, True)

COLORS = {
    "ink": "#222831",
    "muted": "#59616B",
    "grid": "#D9DEE5",
    "accent": "#2B6F77",
    "accent2": "#D17A22",
    "accent3": "#6B7A2A",
    "accent4": "#7A4F9A",
    "bg": "#FFFFFF",
    "soft": "#F5F7F9",
}


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], font, fill, width_px: int, line_gap: int = 4):
    x, y = xy
    words = str(text).split()
    lines = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width_px:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height
    return y


def add_chart_title(draw, title: str, subtitle: str | None, width: int):
    draw.text((50, 34), title, font=FONT_TITLE, fill=hex_to_rgb(COLORS["ink"]))
    if subtitle:
        draw.text((52, 78), subtitle, font=FONT_SMALL, fill=hex_to_rgb(COLORS["muted"]))
    draw.line((50, 110, width - 50, 110), fill=hex_to_rgb(COLORS["grid"]), width=2)


def save_bar_chart(
    path: Path,
    title: str,
    subtitle: str,
    labels: list[str],
    values: list[float],
    value_suffix: str = "",
    colors: list[str] | None = None,
    width: int = 1400,
    height: int = 850,
):
    img = Image.new("RGB", (width, height), COLORS["bg"])
    draw = ImageDraw.Draw(img)
    add_chart_title(draw, title, subtitle, width)
    colors = colors or [COLORS["accent"], COLORS["accent2"], COLORS["accent3"], COLORS["accent4"]]

    left, right, top, bottom = 170, width - 80, 180, height - 150
    max_val = max(values) if values else 1
    max_axis = max(1, math.ceil(max_val / 10) * 10)
    for i in range(6):
        y = bottom - (bottom - top) * i / 5
        draw.line((left, y, right, y), fill=hex_to_rgb(COLORS["grid"]), width=1)
        tick = max_axis * i / 5
        draw.text((70, y - 10), f"{tick:.0f}{value_suffix}", font=FONT_TINY, fill=hex_to_rgb(COLORS["muted"]))

    n = len(labels)
    gap = 36
    bar_w = max(40, int((right - left - gap * (n + 1)) / max(n, 1)))
    for i, (label, val) in enumerate(zip(labels, values)):
        x0 = left + gap + i * (bar_w + gap)
        x1 = x0 + bar_w
        y1 = bottom
        y0 = bottom - (bottom - top) * val / max_axis
        color = hex_to_rgb(colors[i % len(colors)])
        draw.rounded_rectangle((x0, y0, x1, y1), radius=8, fill=color)
        value_text = f"{val:.2f}{value_suffix}" if value_suffix == "%" else f"{val:.0f}{value_suffix}"
        tw = draw.textbbox((0, 0), value_text, font=FONT_BOLD)[2]
        draw.text((x0 + (bar_w - tw) / 2, y0 - 28), value_text, font=FONT_BOLD, fill=hex_to_rgb(COLORS["ink"]))
        draw_wrapped(draw, label, (x0 - 12, bottom + 18), FONT_SMALL, hex_to_rgb(COLORS["ink"]), bar_w + 24)

    draw.line((left, bottom, right, bottom), fill=hex_to_rgb(COLORS["ink"]), width=2)
    img.save(path)


def save_grouped_bar_chart(path: Path, title: str, subtitle: str, data: pd.DataFrame, width=1500, height=900):
    img = Image.new("RGB", (width, height), COLORS["bg"])
    draw = ImageDraw.Draw(img)
    add_chart_title(draw, title, subtitle, width)
    left, right, top, bottom = 160, width - 80, 190, height - 170
    sites = data["site"].tolist()
    series = [c for c in data.columns if c != "site"]
    max_val = max([data[c].max() for c in series] + [1])
    max_axis = 1.0 if max_val <= 1 else max_val

    for i in range(6):
        y = bottom - (bottom - top) * i / 5
        draw.line((left, y, right, y), fill=hex_to_rgb(COLORS["grid"]), width=1)
        draw.text((78, y - 10), f"{max_axis * i / 5:.1f}", font=FONT_TINY, fill=hex_to_rgb(COLORS["muted"]))

    palette = [COLORS["accent"], COLORS["accent2"], COLORS["accent3"], COLORS["accent4"], "#4B8E5F"]
    group_w = (right - left) / len(sites)
    bar_w = max(16, int((group_w - 34) / len(series)))
    for i, site in enumerate(sites):
        base = left + i * group_w + 17
        for j, col in enumerate(series):
            val = float(data.loc[data["site"] == site, col].iloc[0])
            x0 = base + j * bar_w
            x1 = x0 + bar_w - 3
            y0 = bottom - (bottom - top) * val / max_axis
            draw.rectangle((x0, y0, x1, bottom), fill=hex_to_rgb(palette[j % len(palette)]))
        draw.text((base + group_w / 2 - 20, bottom + 18), site, font=FONT_BOLD, fill=hex_to_rgb(COLORS["ink"]))

    legend_start_x = left
    legend_x = legend_start_x
    legend_y = height - 105
    legend_step = 300
    for j, col in enumerate(series):
        if legend_x + legend_step > width - 80:
            legend_x = legend_start_x
            legend_y += 34
        draw.rectangle((legend_x, legend_y, legend_x + 20, legend_y + 20), fill=hex_to_rgb(palette[j % len(palette)]))
        draw.text((legend_x + 28, legend_y - 1), col.replace("_", " ").title(), font=FONT_SMALL, fill=hex_to_rgb(COLORS["ink"]))
        legend_x += legend_step

    draw.line((left, bottom, right, bottom), fill=hex_to_rgb(COLORS["ink"]), width=2)
    img.save(path)


def save_heatmap(path: Path, title: str, subtitle: str, matrix: pd.DataFrame, width=1500, height=950):
    img = Image.new("RGB", (width, height), COLORS["bg"])
    draw = ImageDraw.Draw(img)
    add_chart_title(draw, title, subtitle, width)

    rows = matrix.index.tolist()
    cols = matrix.columns.tolist()
    left, top = 210, 190
    cell_w = min(155, int((width - left - 80) / max(len(cols), 1)))
    cell_h = min(84, int((height - top - 110) / max(len(rows), 1)))

    for j, col in enumerate(cols):
        draw_wrapped(draw, str(col).replace("_", " "), (left + j * cell_w + 8, top - 70), FONT_TINY, hex_to_rgb(COLORS["ink"]), cell_w - 12)
    for i, row in enumerate(rows):
        draw.text((60, top + i * cell_h + cell_h / 2 - 10), str(row), font=FONT_BOLD, fill=hex_to_rgb(COLORS["ink"]))
        for j, col in enumerate(cols):
            val = float(matrix.loc[row, col])
            intensity = int(245 - val * 145)
            color = (intensity, int(252 - val * 120), int(253 - val * 112))
            x0, y0 = left + j * cell_w, top + i * cell_h
            x1, y1 = x0 + cell_w, y0 + cell_h
            draw.rectangle((x0, y0, x1, y1), fill=color, outline=hex_to_rgb(COLORS["grid"]))
            label = f"{val:.2f}"
            bbox = draw.textbbox((0, 0), label, font=FONT_BOLD)
            draw.text((x0 + (cell_w - (bbox[2] - bbox[0])) / 2, y0 + (cell_h - 20) / 2), label, font=FONT_BOLD, fill=hex_to_rgb(COLORS["ink"]))

    img.save(path)


def save_binary_heatmap(path: Path, title: str, subtitle: str, data: pd.DataFrame, width=1800, height=1000):
    matrix = data.set_index("site")
    save_heatmap(path, title, subtitle, matrix, width=width, height=height)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, color="222831", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header_fill="EEF4F5"):
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True


def default_weights(columns: list[str]) -> list[float]:
    weights = []
    for col in columns:
        name = col.lower()
        if any(token in name for token in ["objective", "answer", "limitation", "reason", "needed", "flags", "notes"]):
            weights.append(3.0)
        elif any(token in name for token in ["percent", "count", "site", "value", "result", "growth", "recorded"]):
            weights.append(1.0)
        else:
            weights.append(1.4)
    return weights


def add_df_table(
    doc: Document,
    df: pd.DataFrame,
    columns: list[str] | None = None,
    max_rows: int | None = None,
    weights: list[float] | None = None,
    font_size: int = 8,
):
    view = df.copy()
    if columns:
        view = view[columns]
    if max_rows:
        view = view.head(max_rows)
    table = doc.add_table(rows=1, cols=len(view.columns))
    for i, col in enumerate(view.columns):
        set_cell_text(table.rows[0].cells[i], col.replace("_", " ").title(), bold=True, size=font_size)
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(view.columns):
            set_cell_text(cells[i], row[col], size=font_size)
    style_table(table)
    col_weights = weights or default_weights(list(view.columns))
    widths = column_widths_from_weights(col_weights, CONTENT_WIDTH_DXA)
    apply_table_geometry(table, widths, table_width_dxa=CONTENT_WIDTH_DXA)
    return table


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 97, 107)


def add_note(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_cell_shading(table.cell(0, 0), "F5F7F9")
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(34, 40, 49)
    style_table(table, header_fill="F5F7F9")
    apply_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.color.rgb = RGBColor(34, 40, 49)
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(34, 40, 49)
    return p


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True

    header = section.header
    p = header.paragraphs[0]
    p.text = "Statistical Report | Poso Bacterial Profile Computations"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(89, 97, 107)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.text = "Prepared by Engr. Jamie Eduardo Rosal, MSCpE"
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(89, 97, 107)

    return doc


def main():
    df = pd.read_csv(ROOT / "poso_results_encoded.csv")
    recorded = df[df["results_recorded"] == "Yes"].copy()
    audit = pd.read_csv(ROOT / "computed_outputs/audit.csv")
    objective = pd.read_csv(ROOT / "computed_outputs/objective_coverage.csv")
    growth = pd.read_csv(ROOT / "computed_outputs/growth_summary.csv")
    gram = pd.read_csv(ROOT / "computed_outputs/gram_by_site_counts.csv")
    media = pd.read_csv(ROOT / "computed_outputs/media_summary.csv")
    biochemical = pd.read_csv(ROOT / "computed_outputs/biochemical_summary.csv")
    site_binary = pd.read_csv(ROOT / "computed_outputs/site_binary_profile.csv")
    similarity = pd.read_csv(ROOT / "computed_outputs/site_similarity_jaccard.csv").set_index("site")
    indicators = pd.read_csv(ROOT / "computed_outputs/indicator_flags.csv")
    anova = pd.read_csv(ROOT / "computed_outputs/anova_suitability_check.csv")

    gram_summary = recorded["gram_reaction"].value_counts().rename_axis("gram_reaction").reset_index(name="count")
    gram_summary["percent"] = (gram_summary["count"] / len(recorded) * 100).round(2)

    growth_fig = FIG / "figure_1_growth_status.png"
    gram_fig = FIG / "figure_2_gram_reaction_distribution.png"
    site_features_fig = FIG / "figure_3_site_feature_rates.png"
    similarity_fig = FIG / "figure_4_site_profile_similarity.png"
    biochemical_fig = FIG / "figure_5_selected_biochemical_positive_rates.png"

    save_bar_chart(
        growth_fig,
        "Recorded Bacterial Growth Status",
        "Encoded samples from the research results sheet (n = 18).",
        growth["nutrient_growth"].tolist(),
        growth["percent"].tolist(),
        "%",
        [COLORS["accent"], COLORS["accent2"]],
    )
    save_bar_chart(
        gram_fig,
        "Gram Reaction Distribution",
        "Recorded samples only (n = 15).",
        gram_summary["gram_reaction"].tolist(),
        gram_summary["percent"].tolist(),
        "%",
        [COLORS["accent"], COLORS["accent2"], COLORS["accent4"]],
    )
    selected_features = site_binary[
        ["site", "gram_negative", "emb_metallic_sheen", "motile", "catalase_positive", "oxidase_positive"]
    ].copy()
    save_grouped_bar_chart(
        site_features_fig,
        "Selected Bacterial Profile Features by Site",
        "Values represent the proportion of recorded replicates per site showing each feature.",
        selected_features,
    )
    save_heatmap(
        similarity_fig,
        "Jaccard Similarity of Site Profiles",
        "Higher values indicate more overlap in encoded binary bacterial profile features.",
        similarity,
        width=1300,
        height=850,
    )
    positive = []
    for test, result_label in [
        ("sim_motility", "Positive"),
        ("mr", "Positive"),
        ("catalase", "Positive"),
        ("coagulase", "Positive"),
        ("oxidase", "Positive"),
        ("h2s", "Positive"),
    ]:
        count = int((recorded[test] == result_label).sum())
        positive.append({"test": test, "percent": round(count / len(recorded) * 100, 2)})
    pos_df = pd.DataFrame(positive)
    save_bar_chart(
        biochemical_fig,
        "Selected Positive Biochemical Findings",
        "Percentages are based on recorded samples (n = 15).",
        pos_df["test"].str.replace("_", " ").str.title().tolist(),
        pos_df["percent"].tolist(),
        "%",
        [COLORS["accent3"], COLORS["accent"], COLORS["accent2"], COLORS["accent4"]],
    )

    doc = setup_document()

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Detailed Statistical Report and Interpretation")
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Bacterial Isolation and Profiling of Selected Poso Spout Samples")
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(89, 97, 107)

    meta = doc.add_table(rows=4, cols=2)
    meta_data = [
        ("Prepared by", "Engr. Jamie Eduardo Rosal, MSCpE"),
        ("Analysis platform", "Python computations in Jupyter Notebook"),
        ("Primary dataset", "Encoded observations from RESEARCH RESULTS.pdf"),
        ("Report date", date.today().strftime("%B %d, %Y")),
    ]
    for i, (k, v) in enumerate(meta_data):
        set_cell_text(meta.cell(i, 0), k, bold=True, size=9)
        set_cell_text(meta.cell(i, 1), v, size=9)
    style_table(meta)
    apply_table_geometry(meta, column_widths_from_weights([1.2, 3.8], CONTENT_WIDTH_DXA), table_width_dxa=CONTENT_WIDTH_DXA)

    add_heading(doc, "Abstract", 1)
    add_body(
        doc,
        "This report presents a descriptive statistical analysis of bacterial profiling results from selected poso spout samples. "
        "The analysis was conducted using a structured encoding of laboratory observations, including growth status, Gram reaction, "
        "cultural characteristics, selective and differential media outcomes, and biochemical test results. Because the data are primarily "
        "categorical, the appropriate statistical treatment consists of frequency counts, percentages, cross-tabulations, binary profile "
        "matrices, and site-similarity summaries. The results support descriptive comparison across sampling sites, while species-level "
        "identification and microbial-load inference remain presumptive and limited without confirmatory identification methods or numeric CFU data.",
    )

    add_heading(doc, "Dataset and Analytical Scope", 1)
    add_body(
        doc,
        "The dataset contains 18 encoded sample records representing coded poso samples A-1 to F-3. Of these, 15 samples contained usable "
        "recorded laboratory observations, while three samples under Poso E were blank or missing in the results sheet. All percentages for "
        "morphological, cultural, and biochemical summaries were therefore computed using the 15 recorded samples unless otherwise stated.",
    )
    add_df_table(doc, audit)
    add_caption(doc, "Table 1. Dataset audit and completeness summary.")

    add_note(
        doc,
        "Statistical note: ANOVA is not valid for the current dataset because the available observations are categorical. ANOVA would require "
        "numeric microbial-load measurements such as colony counts, dilution factors, swabbed area, and computed CFU/cm².",
    )

    add_heading(doc, "Objective-Level Answerability", 1)
    add_body(
        doc,
        "The computations directly support the descriptive objectives of the study. The laboratory act of isolating bacteria is not computational, "
        "but the notebook can summarize recorded growth after isolation. Cultural, morphological, biochemical, and between-site profile comparisons "
        "are answerable using the encoded categorical results.",
    )
    add_df_table(
        doc,
        objective,
        ["objective", "answerable_with_notebook", "what_notebook_can_answer", "limitation"],
        weights=[2.7, 1.2, 3.2, 3.2],
        font_size=8,
    )
    add_caption(doc, "Table 2. Objective-level statistical answerability.")

    add_heading(doc, "Results and Interpretation", 1)
    add_heading(doc, "Recorded Growth", 2)
    add_body(
        doc,
        "Bacterial growth was recorded in 15 of 18 encoded samples, corresponding to 83.33% of all sample entries. The remaining 16.67% were "
        "missing because Poso E-1, E-2, and E-3 did not contain recorded observations. Interpreted conservatively, the available results show "
        "that bacterial growth was commonly detected among the sampled poso spouts with completed laboratory records.",
    )
    doc.add_picture(str(growth_fig), width=Inches(6.8))
    add_caption(doc, "Figure 1. Recorded bacterial growth status across encoded samples.")

    add_heading(doc, "Gram Reaction and Morphology", 2)
    add_body(
        doc,
        "Among the 15 recorded samples, Gram-negative organisms predominated. Gram-negative results accounted for 11 samples or 73.33%, "
        "Gram-positive results accounted for 3 samples or 20.00%, and one sample or 6.67% was encoded as Gram-variable. This distribution "
        "indicates that the observed profiles were largely dominated by Gram-negative bacilli or related Gram-negative forms, consistent with "
        "the frequent growth and fermentation observations recorded on selective and differential media.",
    )
    doc.add_picture(str(gram_fig), width=Inches(6.8))
    add_caption(doc, "Figure 2. Distribution of Gram reaction among recorded samples.")

    add_df_table(doc, gram_summary)
    add_caption(doc, "Table 3. Frequency and percentage distribution of Gram reaction.")

    add_heading(doc, "Cultural and Media-Based Characteristics", 2)
    add_body(
        doc,
        "MacConkey, EMB, MSA, SSA, and BAP results were used to summarize cultural characteristics and likely functional groups. Several samples "
        "showed lactose-fermenting or strong lactose-fermenting reactions, while others were repeatedly categorized as non-lactose fermenters. "
        "Metallic green sheen on EMB, where present, was treated only as a presumptive coliform or possible E. coli indicator. The report avoids "
        "definitive species claims because confirmatory identification methods were not provided.",
    )
    add_df_table(doc, media, max_rows=25, weights=[1.5, 2.4, 1.0, 1.0], font_size=8)
    add_caption(doc, "Table 4. Summary of selected media-based cultural characteristics.")

    add_heading(doc, "Biochemical Characteristics", 2)
    add_body(
        doc,
        "The biochemical results show a high frequency of motility-positive observations, with 13 of 15 recorded samples encoded as motile. "
        "H2S production was not observed in the encoded records. MR-positive or weak-positive reactions were present in several samples, while "
        "VP positivity was absent. Catalase, coagulase, and oxidase results varied by site and sample, supporting the conclusion that the poso "
        "spouts did not show a single uniform bacterial profile.",
    )
    doc.add_picture(str(biochemical_fig), width=Inches(6.8))
    add_caption(doc, "Figure 3. Selected positive biochemical findings among recorded samples.")
    add_df_table(doc, biochemical, max_rows=35, weights=[1.5, 2.2, 1.0, 1.0], font_size=8)
    add_caption(doc, "Table 5. Descriptive summary of biochemical test outcomes.")

    add_heading(doc, "Between-Site Profile Comparison", 2)
    add_body(
        doc,
        "Site-level comparison was performed using binary feature rates. Poso A and C were fully Gram-negative among recorded replicates, while "
        "Poso B and D showed mixed Gram-reaction patterns. Poso F showed a mixed profile because one replicate was Gram-variable and other "
        "replicates showed Gram-negative interpretations. Site E was excluded from substantive comparison because all three entries were blank. "
        "The feature-rate matrix indicates that motility was common across A, C, and F, while catalase positivity was prominent in B, D, and F. "
        "Oxidase positivity was strongest in C and F based on the encoded records.",
    )
    doc.add_picture(str(site_features_fig), width=Inches(6.9))
    add_caption(doc, "Figure 4. Selected bacterial profile features by site.")
    add_df_table(doc, site_binary, font_size=7)
    add_caption(doc, "Table 6. Site-level binary feature rates.")

    add_heading(doc, "Profile Similarity", 2)
    add_body(
        doc,
        "Jaccard similarity was used as a descriptive measure of overlap among encoded binary site profiles. The highest off-diagonal similarity "
        "was observed between Poso B and Poso F at 0.80, followed by several comparisons involving Poso F and Poso A or C at 0.75. Poso D showed "
        "lower similarity with Poso C at 0.20, indicating a more distinct categorical profile in the encoded dataset. These values should be read "
        "as descriptive similarity indicators rather than inferential tests.",
    )
    doc.add_picture(str(similarity_fig), width=Inches(6.4))
    add_caption(doc, "Figure 5. Jaccard similarity heatmap of site profiles.")

    add_heading(doc, "Presumptive Indicator Findings", 1)
    add_body(
        doc,
        "Several samples contained patterns that may suggest indicator or opportunistic bacterial groups. For example, EMB metallic green sheen "
        "was encoded as a possible E. coli or coliform indicator, while catalase-positive and coagulase-positive profiles were treated as possible "
        "Staphylococcus aureus indicators. These interpretations are intentionally conservative. They should be described as presumptive because "
        "the dataset does not include molecular confirmation, automated identification, MALDI-TOF, API identification, or other validated confirmatory methods.",
    )
    add_df_table(doc, indicators, ["sample_id", "site", "presumptive_flags"], max_rows=20, weights=[1.0, 0.8, 6.5], font_size=8)
    add_caption(doc, "Table 7. Presumptive indicator flags from encoded laboratory observations.")

    add_heading(doc, "Statistical Suitability and Limitations", 1)
    add_df_table(doc, anova, weights=[1.6, 1.0, 3.2, 3.6], font_size=8)
    add_caption(doc, "Table 8. Suitability of planned statistical methods.")
    add_body(
        doc,
        "The principal limitation is the categorical nature of the result sheet. Counts and percentages are appropriate; inferential testing is "
        "restricted by small sample size, incomplete Site E records, and absence of numeric microbial-load data. Any future ANOVA or microbial-load "
        "comparison should add CFU counts, dilution factors, swabbed area, and standardized units. In addition, inconsistent or ambiguous laboratory "
        "interpretations should be validated by qualified microbiology personnel before final species-level discussion.",
    )

    add_heading(doc, "Conclusion", 1)
    add_body(
        doc,
        "The Jupyter-based computations successfully answer the descriptive components of the research objectives. The encoded results demonstrate "
        "frequent bacterial growth among completed samples, predominance of Gram-negative profiles, diverse media and biochemical reactions, and "
        "observable differences among poso sampling sites. The analysis supports a scholarly Chapter 4-style presentation using descriptive statistics, "
        "tables, and figures. However, the study should avoid overstating definitive species identification or microbial-load differences unless additional "
        "confirmatory and quantitative data are provided.",
    )

    add_heading(doc, "Recommended Chapter 4 Statement", 1)
    add_body(
        doc,
        "Based on the encoded laboratory records, bacterial growth was observed in most completed poso spout samples. Descriptive analysis showed "
        "that Gram-negative bacteria represented the largest proportion of recorded isolates, while biochemical and media-based reactions varied across "
        "sites. These findings indicate heterogeneity in bacterial profiles among selected poso spouts. Because the results are categorical and some "
        "site records are incomplete, the analysis is best interpreted descriptively through frequencies, percentages, and site-profile comparisons rather "
        "than through ANOVA or microbial-load inference.",
    )

    doc.save(REPORT_DOCX)
    print(REPORT_DOCX)
    print("Figures:")
    for p in sorted(FIG.glob("*.png")):
        print(p)


if __name__ == "__main__":
    main()
